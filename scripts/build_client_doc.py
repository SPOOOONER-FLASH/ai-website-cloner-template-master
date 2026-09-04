#!/usr/bin/env python
"""Turns a Markdown document into a Word file for the client.

WHY THIS EXISTS. Client instruction, 2026-09-04: "以后给spooner的md 都用doc，不要md."
Markdown renders as raw asterisks and pipes in anything that is not a code editor, and
the runbook is read on a phone next to a server terminal — the one place where a wall of
`#` and `|` is worst.

WHY THE MARKDOWN STAYS. The .docx is the deliverable, not the source. Markdown is what
git can diff, what a review can comment on line by line, and what the next session can
edit without a Word install. So the repository keeps the .md and this produces the .docx
beside it — the same relationship the site has with `out/`.

WHAT IT HANDLES, because the client docs actually use all of it: headings, paragraphs,
bold and inline code, bullet and numbered lists, tables, fenced code blocks, blockquotes
and rules.

TYPOGRAPHY. Word's defaults are Calibri 11 with a blue-grey heading, which reads as a
template. These documents mix Chinese prose with shell commands, so: a CJK-capable face
for text, a monospace face for anything a person will retype, and near-black headings
that step down in size the way the site's own scale does.

Usage:
  python scripts/build_client_doc.py <input.md> [<input.md> …] [--out <dir>]
  python scripts/build_client_doc.py --all          # every doc in the client set
"""

import re
import sys
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

INK = RGBColor(0x11, 0x11, 0x0F)
INK_2 = RGBColor(0x5A, 0x5A, 0x5E)
RULE = RGBColor(0xDE, 0xDD, 0xD8)
CODE_BG = "F5F5F7"

TEXT_FONT = "Microsoft YaHei"
MONO_FONT = "Consolas"

# The client docs that ship to the desktop folder. Kept here so --all cannot drift from
# what is actually delivered.
CLIENT_DOCS = [
    ("docs/collaboration/CLIENT-RUNBOOK.md", "Spooner操作手册.docx"),
    ("docs/research/GEO-AUDIT-REPORT.md", "GEO审计报告.docx"),
    ("docs/research/BACKLINK_DEEPLINKS.md", "外链改深链清单.docx"),
    ("docs/research/MANUAL_SUBMIT_LIST.md", "需要手动提交的网址.docx"),
]


def set_cell_background(cell, hex_colour):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), hex_colour)
    cell._tc.get_or_add_tcPr().append(shading)


def style_run(run, *, mono=False, bold=False, colour=INK, size=10.5):
    run.font.name = MONO_FONT if mono else TEXT_FONT
    # python-docx sets the Latin face only; CJK needs the East-Asian attribute too, or
    # Word silently substitutes a default for every Chinese character in the run.
    run._element.rPr.rFonts.set(qn("w:eastAsia"), TEXT_FONT)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = colour


INLINE = re.compile(r"(\*\*.+?\*\*|`[^`]+`)")


def add_inline(paragraph, text, *, size=10.5, colour=INK):
    """Renders **bold** and `code` inside a paragraph."""
    for piece in INLINE.split(text):
        if not piece:
            continue
        if piece.startswith("**") and piece.endswith("**"):
            style_run(paragraph.add_run(piece[2:-2]), bold=True, size=size, colour=colour)
        elif piece.startswith("`") and piece.endswith("`"):
            run = paragraph.add_run(piece[1:-1])
            style_run(run, mono=True, size=size - 0.5, colour=colour)
        else:
            style_run(paragraph.add_run(piece), size=size, colour=colour)


def add_heading(doc, text, level):
    sizes = {1: 20, 2: 15, 3: 12.5}
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(20 if level > 1 else 0)
    para.paragraph_format.space_after = Pt(6)
    para.paragraph_format.keep_with_next = True
    add_inline(para, text, size=sizes.get(level, 11), colour=INK)
    for run in para.runs:
        run.font.bold = True
    return para


def add_code_block(doc, lines):
    """One shaded, monospaced cell. A command the reader will retype has to be
    unmistakably separate from the prose telling them to retype it."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = table.cell(0, 0)
    set_cell_background(cell, CODE_BG)
    cell.text = ""
    for index, line in enumerate(lines):
        para = cell.paragraphs[0] if index == 0 else cell.add_paragraph()
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.line_spacing = 1.15
        style_run(para.add_run(line), mono=True, size=9.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def add_table(doc, rows):
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for r, row in enumerate(rows):
        for c, text in enumerate(row):
            cell = table.cell(r, c)
            cell.text = ""
            para = cell.paragraphs[0]
            para.paragraph_format.space_after = Pt(2)
            add_inline(para, text, size=9.5, colour=INK if r == 0 else INK_2)
            if r == 0:
                set_cell_background(cell, CODE_BG)
                for run in para.runs:
                    run.font.bold = True
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def split_table_row(line):
    return [c.strip() for c in line.strip().strip("|").split("|")]


def convert(md_path: Path, out_path: Path):
    doc = Document()

    section = doc.sections[0]
    section.left_margin = section.right_margin = Inches(0.85)
    section.top_margin = section.bottom_margin = Inches(0.8)

    normal = doc.styles["Normal"]
    normal.font.name = TEXT_FONT
    normal.font.size = Pt(10.5)
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), TEXT_FONT)

    lines = md_path.read_text(encoding="utf-8").split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        # Fenced code
        if stripped.startswith("```"):
            i += 1
            block = []
            while i < len(lines) and not lines[i].strip().startswith("```"):
                block.append(lines[i])
                i += 1
            i += 1
            if block:
                add_code_block(doc, block)
            continue

        # Table: a header row followed by a separator row of dashes
        if stripped.startswith("|") and i + 1 < len(lines) and re.match(r"^\|[\s:|-]+\|$", lines[i + 1].strip()):
            rows = [split_table_row(stripped)]
            i += 2
            while i < len(lines) and lines[i].strip().startswith("|"):
                rows.append(split_table_row(lines[i].strip()))
                i += 1
            add_table(doc, rows)
            continue

        # Horizontal rule — rendered as space, not a line. Word's borders are heavier than
        # the markdown reads, and these documents use `---` as a section break that the
        # heading below it already announces.
        if re.fullmatch(r"-{3,}|\*{3,}", stripped):
            doc.add_paragraph().paragraph_format.space_after = Pt(10)
            i += 1
            continue

        heading = re.match(r"^(#{1,4})\s+(.*)$", stripped)
        if heading:
            add_heading(doc, heading.group(2), len(heading.group(1)))
            i += 1
            continue

        if stripped.startswith(">"):
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Inches(0.25)
            para.paragraph_format.space_after = Pt(8)
            add_inline(para, stripped.lstrip("> ").strip(), colour=INK_2)
            i += 1
            continue

        bullet = re.match(r"^([-*+])\s+(.*)$", stripped)
        if bullet:
            para = doc.add_paragraph(style="List Bullet")
            para.paragraph_format.space_after = Pt(3)
            add_inline(para, bullet.group(2))
            i += 1
            continue

        numbered = re.match(r"^(\d+)[.)]\s+(.*)$", stripped)
        if numbered:
            para = doc.add_paragraph(style="List Number")
            para.paragraph_format.space_after = Pt(3)
            add_inline(para, numbered.group(2))
            i += 1
            continue

        para = doc.add_paragraph()
        para.paragraph_format.space_after = Pt(8)
        para.paragraph_format.line_spacing = 1.4
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        add_inline(para, stripped)
        i += 1

    out_path.parent.mkdir(parents=True, exist_ok=True)
    try:
        doc.save(out_path)
        return out_path
    except PermissionError:
        # The client reads these in Word, and Word holds an exclusive lock on the open
        # file. Failing the whole run because one document happens to be on screen would
        # mean the correction never reaches them; writing a dated sibling does.
        stamped = out_path.with_name(f"{out_path.stem}-{date.today():%m%d}{out_path.suffix}")
        doc.save(stamped)
        print(f"    (原文件被 Word 占用，已另存为 {stamped.name})")
        return stamped


def main(argv):
    args = argv[1:]
    out_dir = Path("C:/Users/johns/Desktop/hyde")
    if "--out" in args:
        idx = args.index("--out")
        out_dir = Path(args[idx + 1])
        args = args[:idx] + args[idx + 2 :]

    jobs = []
    if "--all" in args:
        jobs = [(Path(src), out_dir / name) for src, name in CLIENT_DOCS]
    else:
        for a in args:
            if a.startswith("--"):
                continue
            src = Path(a)
            jobs.append((src, out_dir / (src.stem + ".docx")))

    if not jobs:
        print(__doc__)
        return 1

    for src, dst in jobs:
        if not src.exists():
            print(f"  missing: {src}")
            continue
        convert(src, dst)
        print(f"  {src}  ->  {dst}  ({dst.stat().st_size // 1024} KB)")
    return 0


if __name__ == "__main__":
    sys.exit(main(sys.argv))
